#!/usr/bin/env python3
"""
NG911 Filing Tracker - Data Fetcher
====================================
Pulls data from two FCC sources:
  1. ECFS API (docket 25-143) — who has filed NG911 requests
  2. FCC Master PSAP Registry (Socrata/SODA API) — complete PSAP universe

Outputs:
  - ecfs_filings.json        (raw ECFS filing metadata)
  - psap_registry.json       (full PSAP registry)
  - ng911_tracker_data.json  (merged tracker dataset for the dashboard)

Usage:
  python3 fetch_ng911_data.py --api-key YOUR_ECFS_API_KEY
"""

import argparse
import json
import os
import re
import sys
import tempfile
import time
from datetime import datetime
from urllib.request import urlopen, Request
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import docx
    from lxml import etree
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pytesseract
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# ──────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────
ECFS_BASE = "https://publicapi.fcc.gov/ecfs/filings"
PSAP_SODA_BASE = "https://opendata.fcc.gov/resource/dpq5-ta9j.json"
ECFS_DOCKET = "25-143"
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# US state abbreviations for extraction from filer names
US_STATES = {
    'alabama':'AL','alaska':'AK','arizona':'AZ','arkansas':'AR','california':'CA',
    'colorado':'CO','connecticut':'CT','delaware':'DE','florida':'FL','georgia':'GA',
    'hawaii':'HI','idaho':'ID','illinois':'IL','indiana':'IN','iowa':'IA',
    'kansas':'KS','kentucky':'KY','louisiana':'LA','maine':'ME','maryland':'MD',
    'massachusetts':'MA','michigan':'MI','minnesota':'MN','mississippi':'MS',
    'missouri':'MO','montana':'MT','nebraska':'NE','nevada':'NV',
    'new hampshire':'NH','new jersey':'NJ','new mexico':'NM','new york':'NY',
    'north carolina':'NC','north dakota':'ND','ohio':'OH','oklahoma':'OK',
    'oregon':'OR','pennsylvania':'PA','rhode island':'RI','south carolina':'SC',
    'south dakota':'SD','tennessee':'TN','texas':'TX','utah':'UT','vermont':'VT',
    'virginia':'VA','washington':'WA','west virginia':'WV','wisconsin':'WI',
    'wyoming':'WY','district of columbia':'DC','puerto rico':'PR','guam':'GU',
    'virgin islands':'VI','american samoa':'AS','northern mariana islands':'MP',
}
STATE_ABBREVS = set(US_STATES.values())

# Known filer-to-state mappings for filings with no address data in ECFS
FILER_STATE_OVERRIDES = {
    'mclennan county 9-1-1 emergency assistance district': 'TX',
    'bexar metro 9-1-1 network': 'TX',
    'bexar metro 911 network': 'TX',
    'kerr emergency 9-1-1 network': 'TX',
    'city of kilgore': 'TX',
    'texas eastern 9-1-1 network': 'TX',
    'peninsula fiber network, llc': 'MI',
    'alicia atkinson': 'TX',  # Kilgore, TX filing
    'rio grande valley emergency communication district': 'TX',
}

# POC agency overrides — when the form's Q4 agency doesn't reflect the actual ESInet provider
POC_AGENCY_OVERRIDES = {
    'washington military department': 'Comtech',  # Q7 text identifies Comtech as ESInet provider
    'massachusetts state 911 department': 'Comtech',  # MA ESInet provider is Comtech
    'minnesota it services at minnesota department of public safety (mndps),': 'Sinch',  # MN ESInet provider is Sinch
}

# Manual PSAP ID corrections for known data entry errors in FCC filings
PSAP_ID_CORRECTIONS = {
    '4423': '7423',  # Florence County WI filed with wrong ID (4423=Rockingham NC, 7423=Florence WI)
    '4918': '6645',  # Rio Grande Valley TX filed NM Hidalgo 4918; correct TX Hidalgo is 6645
}

# PSAP IDs to ignore — data entry errors, dates, or ZIP codes parsed as PSAP IDs
PSAP_ID_SKIP = {
    '506',    # MN filing listed 506 (Tucson Airport AZ) — no matching MN PSAP
    '2025',   # Date "2025" parsed as PSAP ID (IA Wayne County)
    '2026',   # Date "2026" parsed as PSAP ID (IA Hamilton County)
    '8500',   # ZIP/address from SC filing (NV DPS Elko)
    '29607',  # ZIP code from SC filing
    '29456',  # ZIP code from SC filing
    '05495',  # ZIP/number from SC filing
    '46221',  # ZIP code from IN statewide filing (Fort Wayne area)
    '46204',  # ZIP code from IN statewide filing (Indianapolis)
}

# Known statewide authorities — filings from these cover ALL PSAPs in their state
STATEWIDE_AUTHORITIES = {
    'massachusetts state 911 department': 'MA',
    'ohio department of administrative services 9-1-1 program office': 'OH',
    'ohio department of administrative services 9-1-1 program office ("ohio 911 program office")': 'OH',
    'washington military department - state 911 coordination office': 'WA',
    'state of michigan, state 911 committee': 'MI',
    'indiana 911 board': 'IN',
    'indiana statewide 911 board': 'IN',
    'illinois state police division of statewide 9-1-1': 'IL',
    'minnesota department of public safety (mndps), emergency communications division (ecn)': 'MN',
    'north central texas emergency communications district (nct 9-1-1)': 'TX',
}

# ──────────────────────────────────────────────
# Helper: HTTP GET with retries
# ──────────────────────────────────────────────
def fetch_json(url, retries=3, delay=2):
    for attempt in range(retries):
        try:
            req = Request(url, headers={"User-Agent": "NG911-Tracker/1.0"})
            with urlopen(req, timeout=30) as resp:
                return json.loads(resp.read().decode("utf-8"))
        except HTTPError as e:
            print(f"  HTTP {e.code} on attempt {attempt+1}: {e.reason}")
            if e.code == 429:
                time.sleep(delay * (attempt + 1) * 2)
            elif e.code >= 500:
                time.sleep(delay * (attempt + 1))
            else:
                raise
        except URLError as e:
            print(f"  Network error on attempt {attempt+1}: {e.reason}")
            time.sleep(delay * (attempt + 1))
    raise Exception(f"Failed after {retries} attempts: {url}")


def extract_state_from_name(name):
    """Try to extract a US state from a filer name string."""
    name_lower = name.lower().strip()
    # Check known filer-to-state overrides first
    if name_lower in FILER_STATE_OVERRIDES:
        return FILER_STATE_OVERRIDES[name_lower]
    # Check for full state names (longer names first)
    for state_name, abbrev in sorted(US_STATES.items(), key=lambda x: -len(x[0])):
        if state_name in name_lower:
            return abbrev
    # Check for state abbreviations as standalone words
    for abbrev in STATE_ABBREVS:
        if re.search(r'\b' + abbrev + r'\b', name):
            return abbrev
    return ""


def extract_psap_id_from_name(name):
    """Try to extract PSAP ID from filer name like 'Crawford County Missouri PSAP ID 3789'."""
    match = re.search(r'PSAP\s*ID\s*(\d+)', name, re.IGNORECASE)
    if match:
        return match.group(1)
    return ""


# ──────────────────────────────────────────────
# Step 1: Fetch all ECFS filings for docket 25-143
# ──────────────────────────────────────────────
def fetch_ecfs_filings(api_key, limit_per_page=25):
    """
    Fetch all filings from ECFS for PS Docket 25-143.
    API returns: {"filing": [...], "filingCount": N, "aggregations": {...}}
    Note: the array key is "filing" (singular), not "filings".
    """
    print("\n═══ Step 1: Fetching ECFS Filings (Docket 25-143) ═══")
    all_filings = []
    offset = 0

    while True:
        params = {
            "proceedings.name": ECFS_DOCKET,
            "sort": "date_submission,DESC",
            "limit": limit_per_page,
            "offset": offset,
            "api_key": api_key,
        }
        url = f"{ECFS_BASE}?{urlencode(params)}"
        print(f"  Fetching offset={offset} ...")

        data = fetch_json(url)

        # API uses "filing" (singular) as the array key
        filings = data.get("filing", data.get("filings", []))
        if not filings:
            break

        all_filings.extend(filings)
        total_reported = data.get("filingCount", 0)
        print(f"  Got {len(filings)} filings (total so far: {len(all_filings)}, API total: {total_reported})")

        if total_reported and len(all_filings) >= total_reported:
            break
        if len(filings) < limit_per_page:
            break

        offset += limit_per_page
        time.sleep(0.5)

    print(f"  ✓ Total ECFS filings retrieved: {len(all_filings)}")
    return all_filings


# ──────────────────────────────────────────────
# Step 2: Fetch full PSAP Registry via SODA API
# ──────────────────────────────────────────────
def fetch_psap_registry(limit_per_page=1000):
    print("\n═══ Step 2: Fetching FCC Master PSAP Registry (via API) ═══")
    all_psaps = []
    offset = 0

    while True:
        params = {
            "$limit": limit_per_page,
            "$offset": offset,
            "$order": "state,psap_name",
        }
        url = f"{PSAP_SODA_BASE}?{urlencode(params)}"
        print(f"  Fetching offset={offset} ...")

        psaps = fetch_json(url)
        if not psaps:
            break

        all_psaps.extend(psaps)
        print(f"  Got {len(psaps)} PSAPs (total so far: {len(all_psaps)})")

        if len(psaps) < limit_per_page:
            break

        offset += limit_per_page
        time.sleep(0.3)

    print(f"  ✓ Total PSAPs retrieved: {len(all_psaps)}")
    return all_psaps


def load_psap_registry_from_xlsx(xlsx_path):
    """
    Load PSAP registry from the official FCC Master PSAP Registry XLSX file.
    Download from: https://www.fcc.gov/general/9-1-1-master-psap-registry

    File structure (as of 2026):
      - Rows 0-9: Title, legend, notes
      - Row 10: Header row (PSAP ID, PSAP Name, State, County, City, Type of Change, Comments, Date Last Modified)
      - Row 11+: Data
    """
    try:
        import openpyxl
    except ImportError:
        print("  ✗ openpyxl not installed. Run: pip3 install openpyxl")
        print("  Falling back to Socrata API...")
        return None

    print(f"\n═══ Step 2: Loading PSAP Registry from XLSX ═══")
    print(f"  File: {xlsx_path}")

    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    ws = wb.active

    # Find the header row dynamically
    header_row_idx = None
    headers = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        vals = [str(v).strip().lower() if v else '' for v in row[:10]]
        if 'psap id' in vals:
            header_row_idx = i
            headers = [str(v).strip() if v else '' for v in row[:10]]
            break
        if i > 20:
            break

    if header_row_idx is None:
        print("  ✗ Could not find header row in XLSX")
        wb.close()
        return None

    print(f"  Header at row {header_row_idx}: {headers[:8]}")

    # Map column positions
    col_map = {}
    for idx, h in enumerate(headers):
        hl = h.lower().strip()
        if 'psap id' in hl:
            col_map['psap_id'] = idx
        elif 'psap name' in hl or (hl == 'psap name'):
            col_map['psap_name'] = idx
        elif hl == 'state':
            col_map['state'] = idx
        elif hl == 'county':
            col_map['county'] = idx
        elif hl == 'city':
            col_map['city'] = idx
        elif 'type of change' in hl or 'change' in hl:
            col_map['change_type'] = idx
        elif 'comment' in hl:
            col_map['comments'] = idx
        elif 'date' in hl and 'modified' in hl:
            col_map['date_modified'] = idx

    # Parse data rows
    all_psaps = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i <= header_row_idx:
            continue

        vals = list(row[:10])
        psap_id = vals[col_map.get('psap_id', 0)]
        if psap_id is None:
            continue

        # Determine PSAP type from change_type column
        change_type = str(vals[col_map.get('change_type', 5)] or '').strip()
        comments = str(vals[col_map.get('comments', 6)] or '').strip()

        # Determine if Primary or Secondary
        psap_type = "Primary"
        if change_type == 'S' or 'secondary' in comments.lower():
            psap_type = "Secondary"
        elif change_type == 'O' or 'orphan' in comments.lower():
            psap_type = "Orphaned"

        date_mod = vals[col_map.get('date_modified', 7)]
        if date_mod:
            date_mod = str(date_mod).split(' ')[0]  # Just the date part

        record = {
            "psap_id": str(psap_id).strip(),
            "psap_name": str(vals[col_map.get('psap_name', 1)] or '').strip(),
            "state": str(vals[col_map.get('state', 2)] or '').strip().upper(),
            "county": str(vals[col_map.get('county', 3)] or '').strip(),
            "city": str(vals[col_map.get('city', 4)] or '').strip(),
            "psap_type": psap_type,
            "change_type": change_type,
            "comments": comments,
            "date_modified": date_mod or "",
        }
        all_psaps.append(record)

    wb.close()
    print(f"  ✓ Loaded {len(all_psaps)} PSAPs from XLSX (as of 2/6/2026)")

    # Quick stats
    states = set(p['state'] for p in all_psaps if p['state'])
    primary = sum(1 for p in all_psaps if p['psap_type'] == 'Primary')
    secondary = sum(1 for p in all_psaps if p['psap_type'] == 'Secondary')
    orphaned = sum(1 for p in all_psaps if p['psap_type'] == 'Orphaned')
    print(f"    • {len(states)} states/territories")
    print(f"    • {primary} primary, {secondary} secondary, {orphaned} orphaned")

    return all_psaps


# ──────────────────────────────────────────────
# Step 3: Parse ECFS filings into structured records
# ──────────────────────────────────────────────
def parse_ecfs_filings(raw_filings):
    print("\n═══ Step 3: Parsing ECFS filing metadata ═══")
    parsed = []

    for f in raw_filings:
        # Filer names
        filers = f.get("filers", [])
        filer_names = [fl.get("name", "Unknown") for fl in filers] if filers else ["Unknown"]
        filer_primary = filer_names[0] if filer_names else "Unknown"

        # Extract state from filer name
        filer_state = extract_state_from_name(filer_primary)

        # Extract PSAP ID from filer name
        psap_id = extract_psap_id_from_name(filer_primary)

        # Documents — API uses "src" for URL, "filename" for name
        documents = f.get("documents", [])
        doc_urls = [doc["src"] for doc in documents if doc.get("src")]
        doc_filenames = [doc["filename"] for doc in documents if doc.get("filename")]

        # Filing type from submissiontype object
        sub_type = f.get("submissiontype", {})
        filing_type = sub_type.get("description", "") or sub_type.get("short", "")

        # Filing status
        filing_status = f.get("filingstatus", {}).get("description", "")

        # Infer phase from document filenames or filer name
        all_text = " ".join([filer_primary] + doc_filenames).lower()
        if "phase 2" in all_text or "phase ii" in all_text:
            phase = "Phase 2"
        elif "phase 1" in all_text or "phase i" in all_text:
            phase = "Phase 1"
        else:
            phase = "See PDF"

        record = {
            "filing_id": f.get("id_submission", ""),
            "filer_names": filer_names,
            "filer_primary": filer_primary,
            "filer_state": filer_state,
            "date_submitted": f.get("date_submission", ""),
            "date_received": f.get("date_received", ""),
            "date_disseminated": f.get("date_disseminated", ""),
            "filing_type": filing_type,
            "filing_status": filing_status,
            "document_urls": doc_urls,
            "document_filenames": doc_filenames,
            "document_count": len(doc_urls),
            "phase_requested": phase,
            "psap_id_extracted": psap_id,
        }
        parsed.append(record)

    states_found = sum(1 for f in parsed if f["filer_state"])
    psap_matches = sum(1 for f in parsed if f["psap_id_extracted"])
    print(f"  ✓ Parsed {len(parsed)} filing records")
    print(f"    • State extracted: {states_found}")
    print(f"    • PSAP ID extracted: {psap_matches}")
    if len(parsed) > states_found:
        unmatched = [f["filer_primary"] for f in parsed if not f["filer_state"]][:5]
        print(f"    • Unmatched filers (first 5):")
        for name in unmatched:
            print(f"      - {name}")

    return parsed


# ──────────────────────────────────────────────
# Step 3b: Download PDFs and extract phase/PSAP data
# ──────────────────────────────────────────────
def download_pdf(url, dest_path, retries=2, timeout=90):
    """Download a PDF file from a URL. Handles FCC redirects and large files."""
    import ssl
    from urllib.request import build_opener, HTTPRedirectHandler, HTTPSHandler

    # Build opener that follows redirects and accepts any SSL
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE

    opener = build_opener(
        HTTPSHandler(context=ctx),
        HTTPRedirectHandler()
    )
    opener.addheaders = [
        ("User-Agent", "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"),
        ("Accept", "application/pdf,*/*"),
    ]

    for attempt in range(retries):
        try:
            resp = opener.open(url, timeout=timeout)
            # Stream the download in chunks to avoid memory/timeout issues
            with open(dest_path, "wb") as f:
                while True:
                    chunk = resp.read(65536)  # 64KB chunks
                    if not chunk:
                        break
                    f.write(chunk)
            resp.close()

            # Verify we got a real document (not an HTML error page)
            with open(dest_path, "rb") as f:
                header = f.read(8)
            if b'%PDF' in header:
                pass  # Valid PDF
            elif header[:4] == b'PK\x03\x04':
                pass  # Valid DOCX/ZIP-based Office file
            else:
                print(f"    ✗ Not a PDF/DOCX (got HTML/error page)")
                os.remove(dest_path)
                return False

            return True
        except Exception as e:
            if attempt < retries - 1:
                print(f"    ⟳ Retry {attempt+1}: {str(e)[:60]}")
                time.sleep(3)
            else:
                print(f"    ✗ Failed: {str(e)[:80]}")
                # Clean up partial download
                if os.path.exists(dest_path):
                    os.remove(dest_path)
                return False


def ocr_pdf(pdf_path):
    """OCR a scanned/image PDF using pdf2image + pytesseract."""
    if not HAS_OCR:
        return ""
    try:
        # Use Homebrew paths if standard PATH doesn't include them
        poppler_path = None
        if os.path.isfile("/opt/homebrew/bin/pdftoppm"):
            poppler_path = "/opt/homebrew/bin"
        tesseract_path = "/opt/homebrew/bin/tesseract"
        if os.path.isfile(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path

        kwargs = {}
        if poppler_path:
            kwargs["poppler_path"] = poppler_path
        images = convert_from_path(pdf_path, **kwargs)
        full_text = ""
        for img in images:
            text = pytesseract.image_to_string(img)
            full_text += text + "\n"
        return full_text
    except Exception as e:
        return ""


def extract_phase_from_pdf(pdf_path):
    """
    Extract structured data from an FCC NG911 Valid Request Form PDF.

    Extracts:
      - Q1: 911 Authority name (the actual PSAP/entity, not the person who filed)
      - Q6: Phase selection (Phase 1, Phase 2, or both)
      - Q9: PSAP table (PSAP ID, Name, Phase)

    Based on confirmed pdfplumber text output:
      - Checked boxes render as ☒
      - Unchecked boxes render as ☐
      - Q6 lines: "☒ Phase 1 – Request for OSPs..." / "☐ Phase 2 – Request for OSPs..."
      - Q1 answer appears after "and/or NG911 Phase 2 service (under 47 CFR § 9.31(b))."
      - Q9 PSAP table rows: "PSAP_ID PSAP_Name PHASE_NUM"
    """
    try:
        import logging
        logging.getLogger("pdfminer").setLevel(logging.ERROR)
        with pdfplumber.open(pdf_path) as pdf:
            full_text = ""
            for page in pdf.pages:
                text = page.extract_text() or ""
                full_text += text + "\n"
    except Exception as e:
        return {"phase": "PDF error", "authority_name": "", "poc_agency": "", "psap_ids": [], "psap_table": [], "raw_excerpt": str(e)}

    # OCR fallback for scanned/image PDFs with little extractable text
    if len(full_text.strip()) < 100 and HAS_OCR:
        ocr_text = ocr_pdf(pdf_path)
        if len(ocr_text.strip()) > len(full_text.strip()):
            print(f"    ⟳ OCR fallback: {len(ocr_text.strip())} chars recovered from scanned PDF")
            full_text = ocr_text

    lines = full_text.split('\n')

    # ── Q1: Extract 911 Authority Name ──
    authority_name = ""
    for i, line in enumerate(lines):
        # The answer to Q1 appears right after this line:
        if '9.31(b))' in line and 'full name' not in line.lower():
            # Next non-empty line is the authority name
            for j in range(i+1, min(i+5, len(lines))):
                candidate = lines[j].strip()
                # Skip footnotes, blank lines, and page numbers
                if (candidate and 
                    not candidate.startswith('1 ') and 
                    not candidate.startswith('2 ') and
                    not candidate.isdigit() and
                    'A "911 Authority"' not in candidate and
                    '"Originating service' not in candidate and
                    len(candidate) > 3):
                    authority_name = candidate
                    break
            break

    # Fallback: try Agency field from Q3 contact table
    if not authority_name:
        for i, line in enumerate(lines):
            if line.strip().startswith('Agency') and i > 0:
                # Agency line might be "Agency Santa Rosa County Emergency Comm. Center"
                parts = line.strip().split('Agency', 1)
                if len(parts) > 1 and parts[1].strip():
                    authority_name = parts[1].strip()
                    break

    # ── Q4: Extract Point of Contact Agency (ESInet provider) ──
    # The form has two contact tables with "Agency" fields:
    #   Q3: Person submitting form (their agency = the 911 Authority itself)
    #   Q4: Point of Contact for additional info (their agency = ESInet provider/consultant)
    # We want the Q4 agency. Q4 appears after "Point of Contact Information"
    # and its contact table is the SECOND occurrence of Name/Title/Agency/Phone/Email.
    poc_agency = ""
    in_q4 = False
    agency_occurrences = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if 'point of contact' in stripped.lower() and 'additional information' in stripped.lower():
            in_q4 = True
            continue
        # Collect all Agency lines, we want the one in Q4 section
        if stripped.startswith('Agency') and in_q4:
            parts = stripped.split('Agency', 1)
            if len(parts) > 1 and parts[1].strip():
                poc_agency = parts[1].strip()
                break
        # Also detect if we've left Q4 (hit Q5)
        if in_q4 and '5.' in stripped and 'affected osp' in stripped.lower():
            break

    # ── Q6: Phase Detection ──
    # Look specifically for Q6 "Request Type" checkbox pattern
    # The form text has exactly:
    #   "☒ Phase 1 – Request for OSPs to begin..."  (checked)
    #   "☐ Phase 2 – Request for OSPs to begin..."  (unchecked)
    # These appear in the "6. Request Type" section
    
    phase_1_selected = False
    phase_2_selected = False
    in_q6 = False
    checked_chars = {'☒', '☑', '✓', '✔'}

    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Detect Q6 section
        if '6.' in stripped and 'request type' in stripped.lower():
            in_q6 = True
            continue
        
        # Detect end of Q6 (Q7 starts)
        if in_q6 and ('7.' in stripped and 'phase 1 certification' in stripped.lower()):
            break
        
        if in_q6 or ('phase 1' in stripped.lower() and '–' in stripped) or ('phase 2' in stripped.lower() and '–' in stripped):
            # Build context from prev + current + next lines
            context = stripped
            if i > 0:
                context = lines[i-1].strip() + " " + context
            if i < len(lines) - 1:
                context = context + " " + lines[i+1].strip()
            
            # Check for Phase 1 checkbox line
            if 'phase 1' in stripped.lower() and '–' in stripped:
                for ch in checked_chars:
                    if ch in context:
                        phase_1_selected = True
                        break
            
            # Check for Phase 2 checkbox line
            if 'phase 2' in stripped.lower() and '–' in stripped:
                for ch in checked_chars:
                    if ch in context:
                        phase_2_selected = True
                        break

    # Fallback: scan entire document for the specific pattern
    if not phase_1_selected and not phase_2_selected:
        # Look for "☒ Phase 1" or "Phase 1...☒" anywhere near "Request for OSPs"
        for i, line in enumerate(lines):
            stripped = line.strip()
            # Build wide context (prev + current + next two lines)
            context = stripped
            if i > 0:
                context = lines[i-1].strip() + " " + context
            if i < len(lines) - 1:
                context = context + " " + lines[i+1].strip()
            if i < len(lines) - 2:
                context = context + " " + lines[i+2].strip()
            
            if 'phase 1' in stripped.lower() and ('request for osps' in context.lower() or '9.29(a)' in context):
                for ch in checked_chars:
                    if ch in context:
                        phase_1_selected = True
                        break
            
            if 'phase 2' in stripped.lower() and ('request for osps' in context.lower() or '9.29(b)' in context):
                for ch in checked_chars:
                    if ch in context:
                        phase_2_selected = True
                        break

    # Ultimate fallback: find Q6 block and check for ☒ near Phase text
    if not phase_1_selected and not phase_2_selected:
        # Extract just the Q6 section text (between "Request Type" and "Certifications")
        q6_match = re.search(
            r'(?:6\.\s*Request\s*Type|Request\s*Type|Check one or both)',
            full_text, re.IGNORECASE
        )
        q7_match = re.search(
            r'7\.\s*Phase\s*1\s*Certification',
            full_text, re.IGNORECASE
        )
        if q6_match:
            q6_start = q6_match.start()
            q6_end = q7_match.start() if q7_match else q6_start + 1500
            q6_text = full_text[q6_start:q6_end]
            
            # In Q6 text, check if ☒ appears near Phase 1 or Phase 2
            # Split into Phase 1 block and Phase 2 block
            p1_match = re.search(r'Phase\s*1\s*[–\-]', q6_text)
            p2_match = re.search(r'Phase\s*2\s*[–\-]', q6_text)
            
            if p1_match:
                # Check ~300 chars around Phase 1 mention for ☒
                p1_start = max(0, p1_match.start() - 150)
                p1_end = min(len(q6_text), p1_match.end() + 300)
                p1_block = q6_text[p1_start:p1_end]
                if '☒' in p1_block or '☑' in p1_block or '✓' in p1_block:
                    phase_1_selected = True
            
            if p2_match:
                p2_start = max(0, p2_match.start() - 150)
                p2_end = min(len(q6_text), p2_match.end() + 300)
                p2_block = q6_text[p2_start:p2_end]
                if '☒' in p2_block or '☑' in p2_block or '✓' in p2_block:
                    phase_2_selected = True

    # Determine final phase
    if phase_1_selected and phase_2_selected:
        phase = "Phase 1 & 2"
    elif phase_2_selected:
        phase = "Phase 2"
    elif phase_1_selected:
        phase = "Phase 1"
    else:
        phase = "See PDF"

    # ── Q9: PSAP Table Extraction ──
    # Section starts with "9. Optional - Destination PSAP(s)"
    # Table header: "PSAP ID | PSAP Name | Phase(s) Requested"
    # Data rows like: "1646 Santa Rosa County 1"
    # or multi-line: "1646 Santa Rosa County\nEmergency Comm.\nCenter 1"
    
    psap_table = []
    psap_ids = []
    in_q9 = False
    in_psap_table = False
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Detect Q9 section
        if '9.' in stripped and 'destination psap' in stripped.lower():
            in_q9 = True
            continue
        
        # Detect PSAP table header
        if in_q9 and 'psap' in stripped.lower() and ('name' in stripped.lower() or 'id' in stripped.lower()):
            in_psap_table = True
            continue
        
        # End of Q9 — hit footnotes, privacy act, or next section
        if in_q9 and ('privacy act' in stripped.lower() or 
                       'paperwork reduction' in stripped.lower() or
                       'the commission' in stripped.lower() and 'rules do not require' in stripped.lower()):
            in_psap_table = False
            in_q9 = False
            continue
        
        if in_psap_table and stripped:
            # Skip boilerplate lines
            if stripped.isdigit() and len(stripped) <= 2:
                continue  # page number
            if stripped.startswith('9 ') or stripped.startswith('Phase(s)'):
                continue  # footnote or header
            if 'commission' in stripped.lower() and 'rules' in stripped.lower():
                continue
            if 'or Both)' in stripped:
                continue  # part of table header
            # Skip date fragments (e.g. "May 5," or "January 15, 2025")
            if re.match(r'^(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d', stripped, re.IGNORECASE):
                continue
            # Skip lines that are just "All PSAPs in STATE" (statewide reference, no individual data)
            if re.match(r'^All\s+PSAPs?\s+in\s+', stripped, re.IGNORECASE):
                continue
            # Skip lines about ZIP codes (delivery point info, not PSAP data)
            if 'zip code' in stripped.lower():
                continue
                
            # Try space-separated PSAP IDs first (e.g. "6 6 9 5 1" instead of "66951")
            # Must check before normal match since normal regex would match just the first digit
            space_match = re.match(r'^(\d(?:\s+\d){2,5})\s+(.+?)(?:\s+(1|2|Both|1\s*(?:&|and)\s*2))?\s*$', stripped)
            if space_match and ' ' in space_match.group(1):
                pid = space_match.group(1).replace(' ', '')
                pname = space_match.group(2).strip()
                pphase = space_match.group(3) or ""
                if pid.isdigit() and int(pid) > 50:
                    psap_table.append({
                        "psap_id": pid,
                        "psap_name": pname,
                        "phase": pphase,
                    })
                    if pid not in psap_ids:
                        psap_ids.append(pid)
                continue

            # Clean OCR pipe artifacts: "35EGNC910821 | Harrison County 1" → "35EGNC910821 Harrison County 1"
            cleaned = re.sub(r'\s*\|\s*', ' ', stripped)

            # Try to extract PSAP ID from start of line (1-5 digit number)
            match = re.match(r'^(\d{1,5})\s+(.+?)(?:\s+(1|2|Both|1\s*(?:&|and)\s*2))?\s*$', cleaned)
            if match:
                pid = match.group(1)
                pname = match.group(2).strip()
                pphase = match.group(3) or ""

                # Filter out false positives (page numbers, footnotes, zip codes in wrong section)
                if int(pid) > 50:  # PSAP IDs are typically 100+
                    psap_table.append({
                        "psap_id": pid,
                        "psap_name": pname,
                        "phase": pphase,
                    })
                    if pid not in psap_ids:
                        psap_ids.append(pid)
                continue

            # Non-standard alphanumeric PSAP IDs (e.g. "35EGNC910821 Harrison County 1")
            match2 = re.match(r'^([A-Z0-9]{5,15})\s+(.+?)(?:\s+(1|2|Both))?\s*$', cleaned)
            if match2:
                pid = match2.group(1)
                pname = match2.group(2).strip()
                pphase = match2.group(3) or ""
                if not pid.isalpha() and pid not in ('NG911',) and pid not in psap_ids:
                    psap_table.append({
                        "psap_id": pid,
                        "psap_name": pname,
                        "phase": pphase,
                    })
                    psap_ids.append(pid)
                continue

            # PSAP ID at END of line: "Alcona County 3308" or "Kent County - Grand Rapids PD 3394"
            end_match = re.match(r'^(.+?)\s+(\d{3,5})\s*$', cleaned)
            if end_match:
                pname = end_match.group(1).strip()
                pid = end_match.group(2)
                pid_int = int(pid)
                boilerplate = ('authority', 'registry', 'commission', 'cfr', 'section',
                               'approved by', 'estimated time', 'expires', 'washington')
                if (pid_int > 500 and pid_int != 911 and pid not in psap_ids and
                        not any(b in pname.lower() for b in boilerplate)):
                    psap_table.append({
                        "psap_id": pid,
                        "psap_name": pname,
                        "phase": "",
                    })
                    psap_ids.append(pid)

    # Also extract PSAP IDs from filer name patterns anywhere in doc
    for match in re.finditer(r'PSAP\s*ID\s*[:#]?\s*(\d+)', full_text, re.IGNORECASE):
        pid = match.group(1)
        if pid not in psap_ids:
            psap_ids.append(pid)

    # Scan for attached PSAP list after the form (e.g. Michigan "ALL PSAPs ... see attached")
    # If Q9 referenced an attached list but no PSAPs were found in Q9, scan the doc
    # for "PSAP Name FCC_ID" lines where the ID is at the end, starting after the list header
    has_see_attached = bool(re.search(r'see\s+attached', full_text, re.IGNORECASE))
    has_psap_header = bool(re.search(r'PSAP\s+Name\s+FCC\s*ID', full_text, re.IGNORECASE))
    if (has_see_attached or has_psap_header) and not psap_ids:
        seen_ids = set()
        boilerplate = ('authority', 'registry', 'commission', 'cfr', 'section',
                       'approved by', 'estimated time', 'expires', 'washington',
                       'docket', 'response', 'paperwork', 'privacy', 'date of submission')
        # Only start scanning after the PSAP list header line
        in_psap_list = False
        for line in lines:
            stripped = line.strip()
            if re.match(r'PSAP\s+Name\s+FCC\s*ID', stripped, re.IGNORECASE):
                in_psap_list = True
                continue
            if not in_psap_list:
                continue
            end_match = re.match(r'^(.+?)\s+(\d{3,5})\s*$', stripped)
            if end_match:
                pname = end_match.group(1).strip()
                pid = end_match.group(2)
                pid_int = int(pid)
                if (pid_int > 500 and pid_int != 911 and pid not in seen_ids and
                        not any(b in pname.lower() for b in boilerplate) and
                        len(pname) > 3):
                    seen_ids.add(pid)
                    psap_table.append({
                        "psap_id": pid,
                        "psap_name": pname,
                        "phase": "",
                    })
                    psap_ids.append(pid)

    # Detect "All PSAPs in [STATE]" statewide filings
    statewide_state = ""
    all_psap_match = re.search(r'all\s+psaps?\s+in\s+(\w[\w\s]*?)(?:\s+(?:\(|phase|1|2|both))', full_text, re.IGNORECASE)
    if all_psap_match:
        state_text = all_psap_match.group(1).strip()
        statewide_state = extract_state_from_name(state_text)

    return {
        "phase": phase,
        "authority_name": authority_name,
        "poc_agency": poc_agency,
        "psap_ids": psap_ids,
        "psap_table": psap_table,
        "statewide_state": statewide_state,
        "raw_excerpt": "",
    }


def extract_phase_from_docx(docx_path):
    """
    Extract structured data from an FCC NG911 Valid Request Form in .docx format.
    Same form as the PDF version but submitted as a Word document.
    
    Table layout (confirmed from real filings):
      Table 0: Q1 - Authority name (1x1)
      Table 1: Q2 - Date of submission (1x1)
      Table 2: Q3 - Contact info (5x2: Name/Title/Agency/Phone/Email)
      Table 3: Q4 - POC checkboxes (2x2)
      Table 4: Q4 - POC contact info (5x2: Name/Title/Agency/Phone/Email)
      Table 5: Q5 - Affected OSPs (3x2)
      Table 6: Q6 - Phase checkboxes (2x2)
      Table 7: Q7 - Phase 1 certifications
      Table 8: Q8 - Phase 2 certifications
      Table 9: Q9 - PSAP table (Nx3: ID/Name/Phase)
    
    Checkboxes use Word form fields (w:ffData/w:checkBox) with <w:checked/> when checked.
    """
    try:
        document = docx.Document(docx_path)
    except Exception as e:
        return {"phase": "DOCX error", "authority_name": "", "poc_agency": "",
                "psap_ids": [], "psap_table": [], "raw_excerpt": str(e)}

    tables = document.tables
    
    # ── Q1: Authority Name ──
    authority_name = ""
    if len(tables) > 0 and len(tables[0].rows) > 0:
        authority_name = tables[0].rows[0].cells[0].text.strip()

    # ── Q3: Submitter Agency (the 911 Authority itself) ──
    # Fallback for authority_name if Q1 table is empty
    q3_agency = ""
    if len(tables) > 2:
        for row in tables[2].rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] == 'Agency':
                q3_agency = cells[1]
                break
    if not authority_name and q3_agency:
        authority_name = q3_agency

    # ── Q4: POC Agency (ESInet provider) ──
    poc_agency = ""
    if len(tables) > 4:
        for row in tables[4].rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] == 'Agency':
                poc_agency = cells[1]
                break

    # ── Q6: Phase Selection ──
    # Check Word form field checkboxes in table 6
    phase_1_selected = False
    phase_2_selected = False
    
    if len(tables) > 6:
        phase_table = tables[6]
        for r, row in enumerate(phase_table.rows):
            # Check if this row's text mentions Phase 1 or Phase 2
            row_text = " ".join(c.text for c in row.cells).lower()
            
            # Check the checkbox cell (column 0) for w:checked in XML
            checkbox_cell = row.cells[0]
            cell_xml = etree.tostring(checkbox_cell._tc).decode()
            is_checked = '<w:checked/>' in cell_xml or '<w:checked ' in cell_xml
            
            # Also check for Unicode checkbox chars in cell text
            cell_text = checkbox_cell.text
            if any(ch in cell_text for ch in '☒☑✓✔'):
                is_checked = True
            
            if 'phase 1' in row_text and is_checked:
                phase_1_selected = True
            if 'phase 2' in row_text and is_checked:
                phase_2_selected = True

    # Fallback: check all tables for phase checkbox patterns
    if not phase_1_selected and not phase_2_selected:
        for table in tables:
            for row in table.rows:
                row_text = " ".join(c.text for c in row.cells).lower()
                if 'phase' not in row_text or 'request for osps' not in row_text:
                    continue
                # Check first cell for checkbox
                cell_xml = etree.tostring(row.cells[0]._tc).decode()
                is_checked = '<w:checked/>' in cell_xml or '<w:checked ' in cell_xml
                cell_text = row.cells[0].text
                if any(ch in cell_text for ch in '☒☑✓✔'):
                    is_checked = True
                
                if 'phase 1' in row_text and is_checked:
                    phase_1_selected = True
                if 'phase 2' in row_text and is_checked:
                    phase_2_selected = True

    if phase_1_selected and phase_2_selected:
        phase = "Phase 1 & 2"
    elif phase_2_selected:
        phase = "Phase 2"
    elif phase_1_selected:
        phase = "Phase 1"
    else:
        phase = "See PDF"

    # ── Q9: PSAP Table ──
    psap_table = []
    psap_ids = []
    
    if len(tables) > 9:
        ptable = tables[9]
        for r, row in enumerate(ptable.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 3:
                pid = cells[0].replace('\n', ' ').strip()
                pname = cells[1].replace('\n', ' ').strip()
                pphase = cells[2].replace('\n', ' ').strip()
                
                # Skip header row and empty rows
                if pid.lower().startswith('psap') or not pid or not pid[0].isdigit():
                    continue
                
                psap_table.append({
                    "psap_id": pid,
                    "psap_name": pname,
                    "phase": pphase,
                })
                if pid not in psap_ids:
                    psap_ids.append(pid)
    
    # Also scan all tables for PSAP-ID-like patterns and "All PSAPs" statewide text
    statewide_state = ""
    for table in tables:
        for row in table.rows:
            full_text = " ".join(c.text for c in row.cells)
            for match in re.finditer(r'PSAP\s*ID\s*[:#]?\s*(\d+)', full_text, re.IGNORECASE):
                pid = match.group(1)
                if pid not in psap_ids:
                    psap_ids.append(pid)
            # Detect "All PSAPs in [STATE]"
            all_match = re.search(r'all\s+psaps?\s+in\s+(\w[\w\s]*?)(?:\s|$)', full_text, re.IGNORECASE)
            if all_match and not statewide_state:
                state_text = all_match.group(1).strip()
                statewide_state = extract_state_from_name(state_text)

    return {
        "phase": phase,
        "authority_name": authority_name,
        "poc_agency": poc_agency,
        "psap_ids": psap_ids,
        "psap_table": psap_table,
        "statewide_state": statewide_state,
        "raw_excerpt": "",
    }


def extract_psap_attachment(file_path):
    """
    Extract PSAP IDs and names from non-standard attachment PDFs.
    Handles: standalone PSAP list PDFs, attachment pages within filings,
    multi-page tables, and various formatting styles.
    
    Returns: {"psap_ids": [...], "psap_table": [...]}
    """
    if not HAS_PDFPLUMBER:
        return {"psap_ids": [], "psap_table": []}
    
    try:
        with open(file_path, 'rb') as f:
            header = f.read(4)
    except:
        return {"psap_ids": [], "psap_table": []}

    # If this is a .docx disguised as .pdf, scan ALL tables for PSAP data
    if header[:4] == b'PK\x03\x04':
        if HAS_DOCX:
            try:
                document = docx.Document(file_path)
            except:
                return {"psap_ids": [], "psap_table": []}
            psap_ids = []
            psap_table = []
            for table in document.tables:
                if len(table.rows) < 2:
                    continue
                # Check if this table has PSAP-like columns
                header_cells = [c.text.strip().lower().replace('\n', ' ') for c in table.rows[0].cells]
                id_col = None
                name_col = None
                phase_col = None
                for ci, cell in enumerate(header_cells):
                    if 'psap' in cell and 'id' in cell:
                        id_col = ci
                    elif 'psap' in cell and 'name' in cell:
                        name_col = ci
                    elif 'phase' in cell:
                        phase_col = ci
                # If no PSAP header, check if first data row starts with a number
                if id_col is None and len(table.rows) > 1:
                    first_data = [c.text.strip() for c in table.rows[1].cells]
                    if first_data and first_data[0] and re.match(r'^\d{2,5}$', first_data[0]):
                        id_col = 0
                        name_col = 1 if len(first_data) > 1 else None
                        phase_col = 2 if len(first_data) > 2 else None
                if id_col is None:
                    continue
                for row in table.rows[1:]:
                    cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
                    if len(cells) <= id_col:
                        continue
                    pid = cells[id_col]
                    pname = cells[name_col] if name_col is not None and len(cells) > name_col else ''
                    pphase = cells[phase_col] if phase_col is not None and len(cells) > phase_col else ''
                    if not pid or not pid[0].isdigit():
                        continue
                    if re.match(r'^\d{2,5}$', pid) and int(pid) > 50 and pid not in psap_ids:
                        psap_table.append({"psap_id": pid, "psap_name": pname, "phase": pphase})
                        psap_ids.append(pid)
            if psap_table:
                print(f"    ✓ Attachment (docx) parser found {len(psap_table)} PSAPs")
            return {"psap_ids": psap_ids, "psap_table": psap_table}
        return {"psap_ids": [], "psap_table": []}

    if header != b'%PDF':
        return {"psap_ids": [], "psap_table": []}

    psap_table = []
    psap_ids = []
    seen_ids = set()

    try:
        pdf = pdfplumber.open(file_path)
    except:
        return {"psap_ids": [], "psap_table": []}
    
    try:
        for page in pdf.pages:
            # ── Method 1: pdfplumber table extraction ──
            try:
                tables = page.extract_tables() or []
                for table in tables:
                    if not table or len(table) < 2:
                        continue
                    
                    # Find which column is PSAP ID and which is PSAP Name
                    header_row = table[0]
                    if not header_row:
                        continue
                    
                    id_col = None
                    name_col = None
                    phase_col = None
                    
                    for ci, cell in enumerate(header_row):
                        if not cell:
                            continue
                        cell_lower = cell.lower().replace('\n', ' ')
                        if 'psap' in cell_lower and 'id' in cell_lower:
                            id_col = ci
                        elif 'psap' in cell_lower and ('name' in cell_lower or id_col is not None):
                            name_col = ci
                        elif 'phase' in cell_lower:
                            phase_col = ci
                    
                    # If no header detected, check if first data row starts with a number
                    if id_col is None and len(table) > 1:
                        first_data = table[1]
                        if first_data and first_data[0] and re.match(r'^\d{2,5}$', str(first_data[0]).strip()):
                            id_col = 0
                            name_col = 1 if len(first_data) > 1 else None
                            phase_col = 2 if len(first_data) > 2 else None
                    
                    if id_col is None:
                        continue
                    
                    # Process data rows (skip header)
                    start_row = 1 if any(header_row) and not re.match(r'^\d{2,5}$', str(header_row[0] or '').strip()) else 0
                    
                    for row in table[start_row:]:
                        if not row or len(row) <= id_col:
                            continue
                        
                        raw_id = str(row[id_col] or '').strip()
                        raw_name = str(row[name_col] or '').strip() if name_col is not None and len(row) > name_col else ''
                        raw_phase = str(row[phase_col] or '').strip() if phase_col is not None and len(row) > phase_col else ''
                        
                        # Handle multi-line cells (pdfplumber joins with \n)
                        id_lines = raw_id.split('\n')
                        name_lines = raw_name.split('\n')
                        phase_lines = raw_phase.split('\n')
                        
                        max_lines = max(len(id_lines), len(name_lines))
                        for li in range(max_lines):
                            pid = id_lines[li].strip() if li < len(id_lines) else ''
                            pname = name_lines[li].strip() if li < len(name_lines) else ''
                            pphase = phase_lines[li].strip() if li < len(phase_lines) else ''
                            
                            # Collapse space-separated digits (e.g. "6 6 9 5 1" -> "66951")
                            if re.match(r'^\d(\s+\d){2,5}$', pid):
                                pid = pid.replace(' ', '')
                            # Validate: PSAP ID should be numeric 3-5 digits or alphanumeric pattern
                            if re.match(r'^\d{3,5}$', pid):
                                pid_int = int(pid)
                                # Filter false positives: real PSAP IDs are typically 500+
                                # Also exclude common false positives like 911
                                if pid_int > 500 and pid_int != 911 and pid not in seen_ids:
                                    seen_ids.add(pid)
                                    psap_table.append({
                                        "psap_id": pid,
                                        "psap_name": pname,
                                        "phase": pphase,
                                    })
                                    psap_ids.append(pid)
                            elif re.match(r'^[A-Z0-9]{5,15}$', pid):
                                # Non-standard IDs like 35EGNC910821
                                # Skip common false positives (NG911, etc.)
                                if pid not in seen_ids and pid not in ('NG911',) and not pid.isalpha():
                                    seen_ids.add(pid)
                                    psap_table.append({
                                        "psap_id": pid,
                                        "psap_name": pname,
                                        "phase": pphase,
                                    })
                                    psap_ids.append(pid)
            except Exception as e:
                pass  # table extraction failed, fall through to text method
            
            # ── Method 2: Text-based regex extraction ──
            # For pages where table extraction fails or misses rows
            try:
                text = page.extract_text() or ''
                if not text:
                    continue
                
                for line in text.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Pattern: space-separated PSAP IDs "6 6 9 5 1 Kerr County..."
                    # Check before normal match since normal regex matches just the first digit
                    space_match = re.match(r'^(\d(?:\s+\d){2,5})\s+(.+?)(?:\s+(1|2|Both|1\s*(?:&|and)\s*2))?\s*$', line)
                    if space_match and ' ' in space_match.group(1):
                        pid = space_match.group(1).replace(' ', '')
                        pname = space_match.group(2).strip()
                        pphase = space_match.group(3) or ''
                        if pid.isdigit() and int(pid) > 500 and pid not in seen_ids:
                            seen_ids.add(pid)
                            psap_table.append({
                                "psap_id": pid,
                                "psap_name": pname,
                                "phase": pphase,
                            })
                            psap_ids.append(pid)
                        continue

                    # Pattern: "6013 Abbeville County 9-1-1 1"
                    match = re.match(r'^(\d{2,5})\s+(.+?)(?:\s+(1|2|Both|1\s*(?:&|and)\s*2))?\s*$', line)
                    if match:
                        pid = match.group(1)
                        pname = match.group(2).strip()
                        pphase = match.group(3) or ''

                        pid_int = int(pid)
                        # Filter: real PSAP IDs are 500+, skip boilerplate/zip codes
                        # Also skip 911 itself and if "name" looks like form boilerplate
                        boilerplate = ('authority', 'registry', 'commission', 'cfr', 'section',
                                       'assembly street', 'approved by', 'estimated time')
                        if (pid_int > 500 and pid_int != 911 and pid not in seen_ids and
                                not any(b in pname.lower() for b in boilerplate)):
                            seen_ids.add(pid)
                            psap_table.append({
                                "psap_id": pid,
                                "psap_name": pname,
                                "phase": pphase,
                            })
                            psap_ids.append(pid)
                        continue
                    
                    # Pattern: "35EGNC910821 Harrison County SO 1"
                    match2 = re.match(r'^([A-Z0-9]{5,15})\s+(.+?)(?:\s+(1|2|Both))?\s*$', line)
                    if match2:
                        pid = match2.group(1)
                        pname = match2.group(2).strip()
                        pphase = match2.group(3) or ''

                        if pid not in seen_ids and not pid.isalpha() and pid not in ('NG911',):
                            seen_ids.add(pid)
                            psap_table.append({
                                "psap_id": pid,
                                "psap_name": pname,
                                "phase": pphase,
                            })
                            psap_ids.append(pid)
                        continue

                    # Pattern: PSAP ID at END of line: "Alcona County 3308"
                    end_match = re.match(r'^(.+?)\s+(\d{3,5})\s*$', line)
                    if end_match:
                        pname = end_match.group(1).strip()
                        pid = end_match.group(2)
                        pid_int = int(pid)
                        boilerplate = ('authority', 'registry', 'commission', 'cfr', 'section',
                                       'assembly street', 'approved by', 'estimated time',
                                       'expires', 'washington', 'docket', 'response')
                        if (pid_int > 500 and pid_int != 911 and pid not in seen_ids and
                                not any(b in pname.lower() for b in boilerplate)):
                            seen_ids.add(pid)
                            psap_table.append({
                                "psap_id": pid,
                                "psap_name": pname,
                                "phase": "",
                            })
                            psap_ids.append(pid)
            except:
                pass
    finally:
        pdf.close()

    # OCR fallback: if pdfplumber found nothing, try OCR
    if not psap_table and HAS_OCR:
        ocr_text = ocr_pdf(file_path)
        if ocr_text.strip():
            for line in ocr_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                match = re.match(r'^(\d{2,5})\s+(.+?)(?:\s+(1|2|Both|1\s*(?:&|and)\s*2))?\s*$', line)
                if match:
                    pid = match.group(1)
                    pname = match.group(2).strip()
                    pphase = match.group(3) or ''
                    pid_int = int(pid)
                    if pid_int > 500 and pid_int != 911 and pid not in seen_ids:
                        seen_ids.add(pid)
                        psap_table.append({"psap_id": pid, "psap_name": pname, "phase": pphase})
                        psap_ids.append(pid)
            if psap_table:
                print(f"    ✓ Attachment OCR found {len(psap_table)} PSAPs")

    if psap_table:
        print(f"    ✓ Attachment parser found {len(psap_table)} PSAPs")

    return {"psap_ids": psap_ids, "psap_table": psap_table}


def extract_psap_from_xlsx(file_path):
    """
    Extract PSAP IDs and names from XLSX attachment files.
    These are supplemental PSAP list spreadsheets (e.g. Kansas, GCRPC).
    Returns same format as extract_from_document for merging.
    """
    psap_table = []
    psap_ids = []
    authority_name = ""
    
    try:
        import openpyxl
    except ImportError:
        print(f"    ⚠ openpyxl not installed — cannot parse XLSX")
        return {"phase": "See PDF", "authority_name": "", "poc_agency": "",
                "psap_ids": [], "psap_table": [], "raw_excerpt": "openpyxl not installed"}
    
    # openpyxl validates file extension — if file doesn't end in .xlsx,
    # copy/symlink it with proper extension (common when downloader saves as .pdf)
    actual_path = file_path
    tmp_path = None
    if not file_path.lower().endswith(('.xlsx', '.xlsm', '.xltx', '.xltm')):
        tmp_path = file_path + ".xlsx"
        try:
            import shutil
            shutil.copy2(file_path, tmp_path)
            actual_path = tmp_path
        except:
            pass
    
    try:
        wb = openpyxl.load_workbook(actual_path, read_only=True, data_only=True)
    except Exception as e:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        return {"phase": "See PDF", "authority_name": "", "poc_agency": "",
                "psap_ids": [], "psap_table": [], "raw_excerpt": f"XLSX error: {e}"}
    
    seen_ids = set()
    
    for ws in wb.worksheets:
        id_col = None
        name_col = None
        phase_col = None
        header_row_idx = None
        
        for row_idx, row in enumerate(ws.iter_rows(values_only=False)):
            vals = [str(c.value).strip() if c.value else '' for c in row]
            
            # Look for authority name
            for i, v in enumerate(vals):
                if '911 authority' in v.lower() and i + 1 < len(vals) and vals[i + 1]:
                    authority_name = vals[i + 1].strip().rstrip(':')
            
            # Detect header row by looking for PSAP ID column
            vals_lower = [v.lower() for v in vals]
            for i, v in enumerate(vals_lower):
                if 'psap' in v and 'id' in v:
                    id_col = i
                    header_row_idx = row_idx
                elif 'psap' in v and 'name' in v:
                    name_col = i
                elif 'phase' in v:
                    phase_col = i
            
            # If we've found the header, process data rows
            if header_row_idx is not None and row_idx > header_row_idx:
                raw_id = vals[id_col] if id_col is not None and id_col < len(vals) else ''
                raw_name = vals[name_col] if name_col is not None and name_col < len(vals) else ''
                raw_phase = vals[phase_col] if phase_col is not None and phase_col < len(vals) else ''
                
                if not raw_id:
                    continue
                
                # Clean up: ID might be float from Excel (e.g. "7013.0")
                pid = raw_id.split('.')[0].strip()
                
                # Clean phase: might say "Phase 1" instead of just "1"
                pphase = raw_phase.replace('Phase', '').strip()
                
                # Validate numeric PSAP ID
                if re.match(r'^\d{3,5}$', pid) and int(pid) > 500 and int(pid) != 911:
                    if pid not in seen_ids:
                        seen_ids.add(pid)
                        psap_table.append({
                            "psap_id": pid,
                            "psap_name": raw_name,
                            "phase": pphase,
                        })
                        psap_ids.append(pid)
                # Alphanumeric PSAP IDs
                elif re.match(r'^[A-Z0-9]{5,15}$', pid) and pid not in ('NG911',):
                    if pid not in seen_ids:
                        seen_ids.add(pid)
                        psap_table.append({
                            "psap_id": pid,
                            "psap_name": raw_name,
                            "phase": pphase,
                        })
                        psap_ids.append(pid)
    
    wb.close()
    
    # Clean up temp file if we created one
    if tmp_path and os.path.exists(tmp_path):
        os.remove(tmp_path)
    
    if psap_table:
        print(f"    ✓ XLSX parser found {len(psap_table)} PSAPs")
    
    return {
        "phase": "See PDF",  # XLSX won't have the form's phase declaration
        "authority_name": authority_name,
        "poc_agency": "",
        "psap_ids": psap_ids,
        "psap_table": psap_table,
        "raw_excerpt": "",
    }


def extract_from_document(file_path):
    """
    Route to PDF or DOCX extractor based on file content/extension.
    """
    # Check file header to determine type
    try:
        with open(file_path, 'rb') as f:
            header = f.read(8)
    except:
        return {"phase": "File error", "authority_name": "", "poc_agency": "",
                "psap_ids": [], "psap_table": [], "raw_excerpt": "Cannot read file"}

    if header[:4] == b'%PDF':
        if HAS_PDFPLUMBER:
            return extract_phase_from_pdf(file_path)
        else:
            return {"phase": "See PDF", "authority_name": "", "poc_agency": "",
                    "psap_ids": [], "psap_table": [], "raw_excerpt": "pdfplumber not installed"}
    elif header[:4] == b'PK\x03\x04':
        # ZIP-based format (docx, xlsx, etc.) — distinguish by extension or content
        is_xlsx = False
        try:
            import zipfile
            with zipfile.ZipFile(file_path) as zf:
                names = zf.namelist()
                if any('xl/' in n or 'xl\\' in n for n in names):
                    is_xlsx = True
        except:
            pass
        
        if is_xlsx:
            return extract_psap_from_xlsx(file_path)
        elif HAS_DOCX:
            return extract_phase_from_docx(file_path)
        else:
            return {"phase": "See PDF", "authority_name": "", "poc_agency": "",
                    "psap_ids": [], "psap_table": [], "raw_excerpt": "python-docx not installed"}
    else:
        return {"phase": "See PDF", "authority_name": "", "poc_agency": "",
                "psap_ids": [], "psap_table": [], "raw_excerpt": f"Unknown file type: {header[:4]}"}


def enrich_filings_with_pdfs(parsed_filings):
    """
    Download document attachments and extract phase/PSAP/authority data.
    Supports both PDF (pdfplumber) and DOCX (python-docx) formats.
    Updates filings in-place with extracted information.
    """
    if not HAS_PDFPLUMBER and not HAS_DOCX:
        print("\n═══ Step 3b: Document Extraction (SKIPPED — no parsers installed) ═══")
        print("  Install with: pip3 install pdfplumber python-docx")
        return parsed_filings

    # Process ALL filings to get authority names, not just "See PDF" ones
    to_process = [f for f in parsed_filings if f.get("document_urls")]
    if not to_process:
        print("\n═══ Step 3b: Document Extraction (SKIPPED — no documents) ═══")
        return parsed_filings

    libs = []
    if HAS_PDFPLUMBER: libs.append("PDF")
    if HAS_DOCX: libs.append("DOCX")
    print(f"\n═══ Step 3b: Extracting data from {len(to_process)} documents ({'+'.join(libs)}) ═══")

    pdf_dir = os.path.join(OUTPUT_DIR, "pdfs")
    os.makedirs(pdf_dir, exist_ok=True)

    extracted_phase = 0
    extracted_name = 0
    failed = 0

    for i, filing in enumerate(to_process):
        urls = filing.get("document_urls", [])
        if not urls:
            continue

        filer = filing["filer_primary"][:50]
        print(f"  [{i+1}/{len(to_process)}] {filer}...")

        best_result = None
        for doc_idx, url in enumerate(urls):
            # CRITICAL: The API returns /ecfs/document/ URLs (viewer page).
            # The actual PDF download endpoint is /ecfs/documents/ (with 's').
            download_url = url.replace('/ecfs/document/', '/ecfs/documents/')

            # Download document (could be PDF or DOCX)
            safe_name = re.sub(r'[^\w\-.]', '_', filing["filing_id"])
            # Each document gets its own cache file
            doc_path = os.path.join(pdf_dir, f"{safe_name}_doc{doc_idx}.pdf")

            if os.path.exists(doc_path) and os.path.getsize(doc_path) > 100:
                # Use cached file (skip if suspiciously small/corrupt)
                pass
            else:
                if os.path.exists(doc_path):
                    os.remove(doc_path)  # remove corrupt cache
                success = download_pdf(download_url, doc_path)
                if not success:
                    continue
                time.sleep(0.5)  # be polite

            # Extract data with error protection — handles both PDF and DOCX
            try:
                result = extract_from_document(doc_path)
            except Exception as e:
                print(f"    ✗ Doc {doc_idx+1} parse error: {str(e)[:60]}")
                continue

            # Check if this document has useful form data
            has_phase = result["phase"] not in ("See PDF", "PDF error", "DOCX error", "File error")
            has_name = bool(result.get("authority_name"))
            has_psaps = bool(result.get("psap_ids"))
            has_statewide = bool(result.get("statewide_state"))

            if has_phase or has_name or has_psaps or has_statewide:
                # This doc has form data — use it
                if best_result is None:
                    best_result = result
                else:
                    # Merge: prefer whichever has more data
                    if has_phase and best_result["phase"] in ("See PDF", "PDF error"):
                        best_result["phase"] = result["phase"]
                    if has_name and not best_result.get("authority_name"):
                        best_result["authority_name"] = result["authority_name"]
                    if result.get("poc_agency") and not best_result.get("poc_agency"):
                        best_result["poc_agency"] = result["poc_agency"]
                    if has_psaps:
                        existing = best_result.get("psap_ids", [])
                        best_result["psap_ids"] = list(set(existing + result["psap_ids"]))
                    if result.get("psap_table") and not best_result.get("psap_table"):
                        best_result["psap_table"] = result["psap_table"]
                    if has_statewide and not best_result.get("statewide_state"):
                        best_result["statewide_state"] = result["statewide_state"]

        # Apply best result to filing
        if best_result is None:
            failed += 1
            print(f"    → No extractable data from any document")
            continue

        result = best_result

        # Update phase if we got one
        if result["phase"] != "See PDF" and filing["phase_requested"] == "See PDF":
            filing["phase_requested"] = result["phase"]
            filing["phase_source"] = "PDF extraction"
            extracted_phase += 1

        # Update authority name (the real entity name from Q1)
        if result.get("authority_name"):
            filing["authority_name"] = result["authority_name"]
            filing["filer_display"] = result["authority_name"]
            extracted_name += 1
        else:
            filing["filer_display"] = filing["filer_primary"]

        # Store Q4 Point of Contact Agency (ESInet provider)
        if result.get("poc_agency"):
            filing["poc_agency"] = result["poc_agency"]

        # Add PSAP table from Q9
        if result.get("psap_table"):
            filing["psap_table"] = result["psap_table"]

        # ── Attachment PSAP parser ──
        # If the form parser found few/no PSAPs (e.g. "See Attachment"),
        # run the broader attachment parser on ALL documents for this filing.
        form_psap_count = len(result.get("psap_table", [])) + len(result.get("psap_ids", []))
        if form_psap_count < 3:
            for doc_idx, url in enumerate(urls):
                safe_name = re.sub(r'[^\w\-.]', '_', filing["filing_id"])
                doc_path = os.path.join(pdf_dir, f"{safe_name}_doc{doc_idx}.pdf")
                if os.path.exists(doc_path) and os.path.getsize(doc_path) > 100:
                    attach_result = extract_psap_attachment(doc_path)
                    if attach_result["psap_table"]:
                        # Merge attachment PSAPs with any existing ones
                        existing_ids = set(p["psap_id"] for p in filing.get("psap_table", []))
                        for p in attach_result["psap_table"]:
                            if p["psap_id"] not in existing_ids:
                                filing.setdefault("psap_table", []).append(p)
                                existing_ids.add(p["psap_id"])
                        # Also merge PSAP IDs
                        existing_id_list = filing.get("psap_ids_from_pdf", [])
                        new_ids = [pid for pid in attach_result["psap_ids"] if pid not in existing_id_list]
                        filing["psap_ids_from_pdf"] = existing_id_list + new_ids

        # Add any extracted PSAP IDs
        if result["psap_ids"]:
            existing = filing.get("psap_ids_from_pdf", [])
            filing["psap_ids_from_pdf"] = list(set(existing + result["psap_ids"]))
            if not filing["psap_id_extracted"] and result["psap_ids"]:
                filing["psap_id_extracted"] = result["psap_ids"][0]

        # Store statewide flag — filing covers all PSAPs in a state
        if result.get("statewide_state"):
            filing["statewide_state"] = result["statewide_state"]
            if not filing["filer_state"]:
                filing["filer_state"] = result["statewide_state"]

        # Check known statewide authorities if no statewide flag from doc
        if not filing.get("statewide_state"):
            auth_key = (result.get("authority_name") or filing.get("filer_primary", "")).lower().strip()
            if auth_key in STATEWIDE_AUTHORITIES:
                filing["statewide_state"] = STATEWIDE_AUTHORITIES[auth_key]
                if not filing["filer_state"]:
                    filing["filer_state"] = STATEWIDE_AUTHORITIES[auth_key]
            # Also check filer_primary
            filer_key = filing.get("filer_primary", "").lower().strip()
            if filer_key in STATEWIDE_AUTHORITIES and not filing.get("statewide_state"):
                filing["statewide_state"] = STATEWIDE_AUTHORITIES[filer_key]
                if not filing["filer_state"]:
                    filing["filer_state"] = STATEWIDE_AUTHORITIES[filer_key]

        # Also try to extract state from authority name if we don't have one
        if not filing["filer_state"] and result.get("authority_name"):
            state = extract_state_from_name(result["authority_name"])
            if state:
                filing["filer_state"] = state

        info_parts = []
        if result["phase"] != "See PDF":
            info_parts.append(result["phase"])
        if result.get("authority_name"):
            info_parts.append(f'"{result["authority_name"][:40]}"')
        if result.get("poc_agency"):
            info_parts.append(f'Agency: {result["poc_agency"][:30]}')
        if result.get("psap_ids"):
            info_parts.append(f"PSAPs: {','.join(result['psap_ids'][:3])}")
        print(f"    → {' | '.join(info_parts) if info_parts else 'No structured data found'}")

    # For filings without PDFs, set filer_display to filer_primary
    for f in parsed_filings:
        if "filer_display" not in f:
            f["filer_display"] = f["filer_primary"]

    print(f"  ✓ PDF extraction complete:")
    print(f"    • Phase extracted: {extracted_phase}")
    print(f"    • Authority name extracted: {extracted_name}")
    print(f"    • Download failures: {failed}")

    return parsed_filings


# ──────────────────────────────────────────────
# Step 4: Build the merged tracker dataset
# ──────────────────────────────────────────────
def build_tracker(parsed_filings, psap_registry):
    print("\n═══ Step 4: Building merged tracker dataset ═══")

    # Build county FIPS lookup for PSAP registry
    county_fips_lookup = {}
    try:
        import addfips
        af = addfips.AddFIPS()
        fips_matched = 0
        for p in psap_registry:
            state = (p.get("state") or "").strip()
            county = (p.get("county") or "").strip()
            if state and county:
                fips = af.get_county_fips(county, state=state)
                if fips:
                    key = f"{state}|{county}"
                    county_fips_lookup[key] = fips
                    fips_matched += 1
        print(f"  County FIPS (addfips): matched {fips_matched}/{len(psap_registry)} PSAPs to {len(set(county_fips_lookup.values()))} counties")
    except ImportError:
        # Fallback: try loading pre-built lookup file
        lookup_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "county_fips_lookup.json")
        if os.path.exists(lookup_path):
            try:
                with open(lookup_path) as f:
                    county_fips_lookup = json.load(f)
                print(f"  County FIPS (cached file): {len(county_fips_lookup)} county mappings loaded")
            except:
                print("  ⚠ Could not load county_fips_lookup.json")
        else:
            print("  ⚠ addfips not installed and no county_fips_lookup.json found")
            print("    County-level map coloring won't work.")
            print("    Fix: pip3 install addfips  OR  place county_fips_lookup.json alongside this script")

    # Apply POC agency overrides
    filer_agency_overrides = {
        'minnesota department of public safety': 'Sinch',
    }
    for f in parsed_filings:
        agency = f.get("poc_agency", "").strip()
        if agency.lower() in POC_AGENCY_OVERRIDES:
            f["poc_agency"] = POC_AGENCY_OVERRIDES[agency.lower()]
        elif not agency:
            filer = f.get("filer_primary", "").strip().lower()
            for key, val in filer_agency_overrides.items():
                if key in filer:
                    f["poc_agency"] = val
                    break

    # Apply PSAP ID corrections and remove known-bad IDs before building maps
    for f in parsed_filings:
        pid = f.get("psap_id_extracted", "")
        if pid in PSAP_ID_SKIP:
            f["psap_id_extracted"] = ""
        elif pid in PSAP_ID_CORRECTIONS:
            f["psap_id_extracted"] = PSAP_ID_CORRECTIONS[pid]
        for entry in f.get("psap_table", []):
            epid = entry.get("psap_id", "")
            if epid in PSAP_ID_SKIP:
                entry["psap_id"] = ""
            elif epid in PSAP_ID_CORRECTIONS:
                entry["psap_id"] = PSAP_ID_CORRECTIONS[epid]
        f["psap_ids_from_pdf"] = [PSAP_ID_CORRECTIONS.get(p, p) for p in f.get("psap_ids_from_pdf", []) if p not in PSAP_ID_SKIP]

    # Infer filer_state from PSAP IDs when all PSAPs are in the same state
    psap_id_to_state = {}
    for p in psap_registry:
        pid = str(p.get("psap_id", "")).strip()
        st = (p.get("state") or "").upper().strip()
        if pid and st:
            psap_id_to_state[pid] = st

    inferred_count = 0
    for f in parsed_filings:
        if f.get("filer_state", "").strip():
            continue  # already has a state
        # Collect all PSAP IDs from this filing
        all_pids = set()
        pid = f.get("psap_id_extracted", "")
        if pid:
            all_pids.add(pid)
        for entry in f.get("psap_table", []):
            epid = entry.get("psap_id", "")
            if epid:
                all_pids.add(epid)
        for epid in f.get("psap_ids_from_pdf", []):
            if epid:
                all_pids.add(epid)
        if not all_pids:
            continue
        # Look up states for all PSAP IDs
        states = set()
        for pid in all_pids:
            st = psap_id_to_state.get(pid)
            if st:
                states.add(st)
        # Only infer if all PSAPs resolve to the same single state
        if len(states) == 1:
            inferred_state = states.pop()
            f["filer_state"] = inferred_state
            inferred_count += 1
    if inferred_count:
        print(f"  ✓ Inferred filer_state for {inferred_count} filings from PSAP IDs")

    # PSAP-level filing map
    psap_id_filings = {}
    for f in parsed_filings:
        pid = f.get("psap_id_extracted", "")
        if pid:
            psap_id_filings.setdefault(pid, []).append(f)
        # Also map from psap_table entries
        for entry in f.get("psap_table", []):
            epid = entry.get("psap_id", "")
            if epid and epid not in psap_id_filings.get(epid, []):
                psap_id_filings.setdefault(epid, []).append(f)
        # Also map from psap_ids_from_pdf
        for epid in f.get("psap_ids_from_pdf", []):
            if epid not in psap_id_filings.get(epid, []):
                psap_id_filings.setdefault(epid, []).append(f)

    # Expand statewide filings: match to all PSAPs in the state
    statewide_filings = {}
    for f in parsed_filings:
        sw_state = f.get("statewide_state", "")
        if sw_state:
            statewide_filings.setdefault(sw_state.upper(), []).append(f)

    # State-level filing summary
    states_with_filings = {}
    for f in parsed_filings:
        state = f.get("filer_state", "").upper().strip()
        if not state:
            continue
        states_with_filings.setdefault(state, []).append({
            "filing_id": f["filing_id"],
            "filer": f["filer_primary"],
            "date": f["date_submitted"],
            "type": f["filing_type"],
            "phase": f["phase_requested"],
            "psap_id": f.get("psap_id_extracted", ""),
        })

    # PSAP records
    psap_records = []
    for p in psap_registry:
        state = (p.get("state") or "").upper().strip()
        psap_id = str(p.get("psap_id", "")).strip()

        direct_match = psap_id in psap_id_filings
        statewide_match = state in statewide_filings
        state_match = state in states_with_filings

        if direct_match:
            status = "Filed (PSAP match)"
            filing_count = len(psap_id_filings[psap_id])
        elif statewide_match:
            status = "Filed (statewide)"
            filing_count = len(statewide_filings[state])
        elif state_match:
            status = "Filed (state-level)"
            filing_count = len(states_with_filings[state])
        else:
            status = "No filing detected"
            filing_count = 0

        psap_records.append({
            "psap_id": psap_id,
            "psap_name": p.get("psap_name", ""),
            "state": state,
            "county": p.get("county", ""),
            "county_fips": county_fips_lookup.get(f"{state}|{p.get('county','').strip()}", ""),
            "city": p.get("city", ""),
            "psap_type": p.get("psap_type", ""),
            "state_has_filings": state_match or direct_match,
            "direct_psap_match": direct_match,
            "state_filing_count": filing_count,
            "ng911_status": status,
        })

    total_psaps = len(psap_records)
    all_states = set(p["state"] for p in psap_records if p["state"])
    states_not_filed = all_states - set(states_with_filings.keys())
    direct_matches = sum(1 for p in psap_records if p["direct_psap_match"])

    summary = {
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "total_psaps_in_registry": total_psaps,
        "total_ecfs_filings": len(parsed_filings),
        "direct_psap_matches": direct_matches,
        "states_with_filings": sorted(states_with_filings.keys()),
        "states_without_filings": sorted(states_not_filed),
        "state_filing_count": len(states_with_filings),
        "state_no_filing_count": len(states_not_filed),
        "filings_by_state": {
            state: {
                "count": len(filings),
                "latest_date": max((f["date"] for f in filings if f["date"]), default=None),
                "filers": list(set(f["filer"] for f in filings)),
                "phases": list(set(f["phase"] for f in filings)),
                "psap_ids": list(set(f["psap_id"] for f in filings if f["psap_id"])),
            }
            for state, filings in states_with_filings.items()
        },
    }

    tracker = {
        "summary": summary,
        "filings": parsed_filings,
        "psaps": psap_records,
    }

    print(f"  ✓ Tracker built:")
    print(f"    • {total_psaps} PSAPs in registry")
    print(f"    • {len(parsed_filings)} ECFS filings")
    print(f"    • {len(states_with_filings)} states with filings")
    print(f"    • {len(states_not_filed)} states without filings")
    print(f"    • {direct_matches} direct PSAP ID matches")

    return tracker


# ──────────────────────────────────────────────
# Step 5: Save outputs
# ──────────────────────────────────────────────
def save_outputs(raw_filings, psap_registry, tracker):
    print("\n═══ Step 5: Saving output files ═══")

    path1 = os.path.join(OUTPUT_DIR, "ecfs_filings.json")
    with open(path1, "w") as f:
        json.dump(raw_filings, f, indent=2)
    print(f"  ✓ {path1} ({len(raw_filings)} filings)")

    path2 = os.path.join(OUTPUT_DIR, "psap_registry.json")
    with open(path2, "w") as f:
        json.dump(psap_registry, f, indent=2)
    print(f"  ✓ {path2} ({len(psap_registry)} PSAPs)")

    path3 = os.path.join(OUTPUT_DIR, "ng911_tracker_data.json")
    with open(path3, "w") as f:
        json.dump(tracker, f, indent=2)
    print(f"  ✓ {path3} (merged tracker)")

    return path3


# ──────────────────────────────────────────────
# Step 6 (Optional): Fetch Mutual Agreements (Docket 25-145)
# ──────────────────────────────────────────────
ECFS_DOCKET_MA = "25-145"

def fetch_mutual_agreements(api_key, limit_per_page=25):
    """
    Fetch all filings from ECFS for PS Docket 25-145 (mutual agreements).
    Completely separate from 25-143 valid request processing.
    Saves raw filing metadata — no PDF parsing, no PSAP matching.
    """
    print("\n═══ Step 6: Fetching Mutual Agreements (Docket 25-145) ═══")
    all_filings = []
    offset = 0

    while True:
        params = {
            "proceedings.name": ECFS_DOCKET_MA,
            "sort": "date_submission,DESC",
            "limit": limit_per_page,
            "offset": offset,
            "api_key": api_key,
        }
        url = f"{ECFS_BASE}?{urlencode(params)}"
        print(f"  Fetching offset={offset} ...")

        data = fetch_json(url)
        filings = data.get("filing", data.get("filings", []))
        if not filings:
            break

        all_filings.extend(filings)
        total_reported = data.get("filingCount", 0)
        print(f"  Got {len(filings)} filings (total so far: {len(all_filings)}, API total: {total_reported})")

        if total_reported and len(all_filings) >= total_reported:
            break
        if len(filings) < limit_per_page:
            break

        offset += limit_per_page
        time.sleep(0.5)

    print(f"  ✓ Total mutual agreement filings: {len(all_filings)}")

    # Light parsing — just extract the useful metadata
    parsed = []
    for f in all_filings:
        filers = f.get("filers", [])
        filer_names = [fl.get("name", "") for fl in filers]
        filer_primary = filer_names[0] if filer_names else "Unknown"

        # Filing ID — 25-145 uses 'id_submission'
        filing_id = str(f.get("id_submission", "") or f.get("id_long", "") or f.get("id", "") or "")

        # Filing type — 25-145 uses submissiontype.description
        filing_type = ""
        st = f.get("submissiontype")
        if st:
            filing_type = st.get("description", "") or st.get("short", "")
        if not filing_type:
            filing_type = f.get("type_of_filing", "")

        # Authors
        authors = [a.get("name", "") for a in f.get("authors", [])]

        # Get document URLs and titles — 25-145 uses 'src' and 'filename'
        docs = f.get("documents", [])
        doc_info = []
        for d in docs:
            doc_url = d.get("src", "") or d.get("url", "")
            if not doc_url and d.get("id"):
                doc_url = f"https://www.fcc.gov/ecfs/document/{d['id']}/1"
            doc_title = d.get("filename", "") or d.get("title", "") or d.get("description", "")
            doc_info.append({
                "id": d.get("id", ""),
                "title": doc_title,
                "url": doc_url,
            })

        parsed.append({
            "filing_id": filing_id,
            "date_received": f.get("date_received", ""),
            "date_submission": f.get("date_submission", ""),
            "filer": filer_primary,
            "all_filers": filer_names,
            "authors": authors,
            "type_of_filing": filing_type,
            "documents": doc_info,
            "ecfs_url": f"https://www.fcc.gov/ecfs/filing/{filing_id}" if filing_id else "",
        })

    # Save to separate file
    path = os.path.join(OUTPUT_DIR, "mutual_agreements_25145.json")
    with open(path, "w") as fp:
        json.dump({
            "docket": ECFS_DOCKET_MA,
            "fetched_at": datetime.utcnow().isoformat() + "Z",
            "count": len(parsed),
            "filings": parsed,
        }, fp, indent=2)
    print(f"  ✓ Saved to {path}")

    return parsed


# ──────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="NG911 Filing Tracker - Data Fetcher")
    parser.add_argument("--api-key", required=True, help="FCC ECFS API key")
    parser.add_argument("--skip-psap", action="store_true", help="Skip PSAP registry fetch (use cached)")
    parser.add_argument("--skip-ecfs", action="store_true", help="Skip ECFS fetch (use cached)")
    parser.add_argument("--skip-pdf", action="store_true", help="Skip PDF download/extraction")
    parser.add_argument("--skip-ma", action="store_true", help="Skip mutual agreements (25-145) fetch")
    parser.add_argument("--psap-xlsx", default=None, help="Path to FCC Master PSAP Registry XLSX file (recommended, most current data)")
    args = parser.parse_args()

    print("╔══════════════════════════════════════════╗")
    print("║   NG911 FILING TRACKER — DATA FETCHER   ║")
    print("╚══════════════════════════════════════════╝")
    print(f"  Time: {datetime.utcnow().isoformat()}Z")
    print(f"  Docket: PS {ECFS_DOCKET}")

    if args.skip_ecfs and os.path.exists(os.path.join(OUTPUT_DIR, "ecfs_filings.json")):
        print("\n  [Skipping ECFS fetch — using cached data]")
        with open(os.path.join(OUTPUT_DIR, "ecfs_filings.json")) as f:
            raw_filings = json.load(f)
    else:
        raw_filings = fetch_ecfs_filings(args.api_key)

    if args.psap_xlsx:
        psap_registry = load_psap_registry_from_xlsx(args.psap_xlsx)
        if psap_registry is None:
            print("  Falling back to Socrata API...")
            psap_registry = fetch_psap_registry()
    elif args.skip_psap and os.path.exists(os.path.join(OUTPUT_DIR, "psap_registry.json")):
        print("\n  [Skipping PSAP fetch — using cached data]")
        with open(os.path.join(OUTPUT_DIR, "psap_registry.json")) as f:
            psap_registry = json.load(f)
    else:
        psap_registry = fetch_psap_registry()

    parsed_filings = parse_ecfs_filings(raw_filings)

    if not args.skip_pdf:
        parsed_filings = enrich_filings_with_pdfs(parsed_filings)
    else:
        print("\n  [Skipping PDF extraction]")

    tracker = build_tracker(parsed_filings, psap_registry)
    output_path = save_outputs(raw_filings, psap_registry, tracker)

    # Fetch mutual agreements (25-145) — completely separate from 25-143
    if not args.skip_ma:
        try:
            ma_filings = fetch_mutual_agreements(args.api_key)
        except Exception as e:
            print(f"\n  ⚠ Mutual agreements fetch failed: {e}")
            print(f"    (This doesn't affect 25-143 data)")
    else:
        print("\n  [Skipping mutual agreements (25-145) fetch]")

    print("\n══════════════════════════════════════════")
    print(f"  ✅ Done! Load this file in the dashboard:")
    print(f"     {output_path}")
    print("══════════════════════════════════════════\n")


if __name__ == "__main__":
    main()
